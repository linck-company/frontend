from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from io import BytesIO

api_data = {
    "event_title": "FAST WEDNESDAYS",
    "event_nature": "Competition",
    "event_date": "12-03-2024",
    "event_time": "4:00 pm",
    "event_host": "Cubing Club",
    "event_location": "S310",
    "event_number_attendees": {
        "external_attendees": 50,
        "faculty_attendees": 10,
        "phd_attendees": 5,
        "pg_attendees": 15,
        "ug_attendees": 20
    },
    "event_organizer_contact_info": "Cubing Club",
    "event_subject_area": "Cubing",
    "event_resource_person": "NA",
    "event_affiliation": "Cubing Club",
    "event_resource_person_profile": "NA",
    "event_objective": "The Objective of the event is to teach new solving techniques and allow members to practice different cube puzzles (3x3, Pyraminx, mirror cube etc.). Organize friendly competitions, time trials, or team-solving challenges to make it fun and engaging.",
    "event_outcome": "The outcomes of this event are that participants enhanced their solving techniques and speed, Participants learnt new algorithms, tricks, and strategies from each other. Participants experienced a sense of accomplishment through challenges or competitions.",
    "event_flyer_image_url": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
    "event_attendance": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
    "event_certificate_image_url": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
    "event_winner_details": "The winners of the event are the participants who performed the best in the event.",
    "event_recommended_actions": "Participants should practice the techniques they learned in the event and continue to improve their solving skills.",
    "event_expenses": "The event expenses will be covered by the Cubing Club.",
    "event_revenue": "The event revenue will be used to cover the expenses of the event.",
    "event_pictures": {
        "event_picture_1": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
        "event_picture_2": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
        "event_picture_3": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0",
        "event_picture_4": "https://drive.usercontent.google.com/download?id=1J8BesUsIu6c3npyXsqtnM1WlocpvB15H&authuser=0"
    }
}

def insert_image(cell, image_url):
    """ Inserts an image from a URL into a table cell """
    try:
        response = requests.get(image_url)
        if response.status_code == 200:
            image_stream = BytesIO(response.content)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(2))  # Resize as needed
        else:
            cell.text = "[Image Not Found]"
    except Exception:
        cell.text = "[Error Loading Image]"

def add_header(doc, logo_url):
    """ Adds a header with a right-aligned logo """
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Align logo to the right

    try:
        response = requests.get(logo_url)
        if response.status_code == 200:
            image_stream = BytesIO(response.content)
            run = paragraph.add_run()
            run.add_picture(image_stream, width=Inches(1))  # Adjust size as needed
        else:
            paragraph.add_run("[Logo Not Found]")
    except Exception:
        paragraph.add_run("[Error Loading Logo]:")
    
def create_cubing_report(file_path, api_data):
    doc = Document()
    add_header(doc, "https://drive.usercontent.google.com/download?id=1tsJc46se47Hlrxqhkxs0mdP_OUW_04R3&authuser=0")

    def add_heading(text):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)

    def add_text(text):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.add_run(text).font.size = Pt(11)

    add_heading("\t\t\tDirectorate of Student Affairs")

    event_details = [
        ("Title of the Event:", api_data["event_title"]),
        ("Nature of the Event:", api_data["event_nature"]),
        ("Date:", api_data["event_date"]),
        ("Time:", api_data["event_time"]),
        ("Host:", api_data["event_host"]),
        ("Location:", api_data["event_location"])
    ]

    after_location_details = [
        ("Event Organiser Contact info:", api_data["event_organizer_contact_info"]),
        ("Subject Area:", api_data["event_subject_area"]),
        ("Resource Person / Chief Guest:", api_data["event_resource_person"]),
        ("Affiliation:", api_data["event_affiliation"]),
        ("Resource person profile:", api_data["event_resource_person_profile"]),
        ("Objective of the event:", api_data["event_objective"]),
        ("Outcome of the event:", api_data["event_outcome"]),
        ("Event Flyer/brochure:", api_data["event_flyer_image_url"]),
        ("Attendance Sheet:", api_data["event_attendance"]),
        ("Participant’s Feedback:", "Yes"),
        ("Event Certificates provided to the participants:", api_data["event_certificate_image_url"]),
        ("Details of Winner:", api_data["event_winner_details"]),
        ("Recommended actions:", api_data["event_recommended_actions"]),
        ("Expenses incurred:", api_data["event_expenses"]),
        ("Revenue Generated if Any:", api_data["event_revenue"])
    ]

    # Create table with 6 columns (for Number of attendees) and 3 columns (for Event Pictures)
    total_rows = len(event_details) + len(after_location_details) + 4
    table = doc.add_table(rows=total_rows, cols=6)  # 6 columns for attendees

    table.style = 'Table Grid'

    row_idx = 0
    for key, value in event_details:
        table.cell(row_idx, 0).text = key
        table.cell(row_idx, 1).merge(table.cell(row_idx, 5))  
        if "https:" in value.lower():
            insert_image(table.cell(row_idx, 1), value)  
        else:
            table.cell(row_idx, 1).text = value

        row_idx += 1

    # Insert "Number of attendees" row
    table.cell(row_idx, 0).text = "Number of attendees:"
    headers = ["External", "Faculty", "PhD", "PG", "UG"]
    values = [
        str(api_data["event_number_attendees"]["external_attendees"]),
        str(api_data["event_number_attendees"]["faculty_attendees"]),
        str(api_data["event_number_attendees"]["phd_attendees"]),
        str(api_data["event_number_attendees"]["pg_attendees"]),
        str(api_data["event_number_attendees"]["ug_attendees"])
    ]

    for i in range(5):
        table.cell(row_idx, i + 1).text = headers[i]
        table.cell(row_idx + 1, i + 1).text = values[i]

    table.cell(row_idx + 1, 0).merge(table.cell(row_idx, 0))
    row_idx += 2

    # Fill in remaining details
    for key, value in after_location_details:
        table.cell(row_idx, 0).text = key
        table.cell(row_idx, 1).merge(table.cell(row_idx, 5))
        if "https:" in value.lower():
            insert_image(table.cell(row_idx, 1), value)  
        else:
            table.cell(row_idx, 1).text = value
        row_idx += 1

    # Insert "Event Pictures (with Geo-tagging)" row
    table.cell(row_idx, 0).text = "Event Pictures (with Geo-tagging):"
    # Merge extra columns to ensure only 2 columns are used for images
    table.cell(row_idx, 1).merge(table.cell(row_idx, 2))  # Merge columns 1 & 2
    table.cell(row_idx + 1, 1).merge(table.cell(row_idx + 1, 2))  # Merge below row
    table.cell(row_idx, 3).merge(table.cell(row_idx, 5))  # Merge columns 3-5
    table.cell(row_idx + 1, 3).merge(table.cell(row_idx + 1, 5))  # Merge below row

    # Insert 4 images
    picture_urls = list(api_data["event_pictures"].values())
    for i in range(2):
        insert_image(table.cell(row_idx, i * 2 + 1), picture_urls[i])
        insert_image(table.cell(row_idx + 1, i * 2 + 1), picture_urls[i + 2])

    table.cell(row_idx + 1, 0).merge(table.cell(row_idx, 0))

    add_text("\nChecklist to be submitted:")

    checklist_table = doc.add_table(rows=7, cols=3)
    checklist_table.style = 'Table Grid'
    headers = ["SI. No.", "Name of the Document", "Submitted"]
    for i in range(3):
        checklist_table.cell(0, i).text = headers[i]

    checklist_data = [
        ("1", "Event Flyer/brochure", "Yes"),
        ("2", "Event photographs with geotagging", "Yes"),
        ("3", "Budget approval copy", "N/A"),
        ("4", "Proofs of payments received if any", "N/A"),
        ("5", "Event attendance sheet with attendees' signatures", "Yes"),
        ("6", "Sample certificates copies (if any)", "N/A")
    ]

    for i, (num, doc_name, submitted) in enumerate(checklist_data, 1):
        checklist_table.cell(i, 0).text = num
        checklist_table.cell(i, 1).text = doc_name
        checklist_table.cell(i, 2).text = submitted

    add_text("\tPrepared By: S.Pranavi Navya\t\t\t Signature: ")
    add_text("\t\t\t\t\t\t\t Date: 12-02-2025\n\n\n\n\n")

    sign_table = doc.add_table(rows=4, cols=3)
    sign_table.style = 'Table Grid'
    sign_table.cell(0, 0).text = "Department"
    sign_table.cell(0, 1).text = "Verified By"
    sign_table.cell(0, 2).text = "Signature"
    sign_table.cell(1, 0).text = "Department"
    sign_table.cell(1, 1).text = "Faculty Advisor"
    sign_table.cell(1, 2).text = "[sign]"
    sign_table.cell(2, 0).text = "QAR – Office"
    sign_table.cell(2, 1).text = "[name]"
    sign_table.cell(2, 2).text = "[sign]"
    sign_table.cell(3, 0).text = "CFAO – Office"
    sign_table.cell(3, 1).text = "[name]"
    sign_table.cell(3, 2).text = "[sign]"

    doc.save(file_path)
    print(f"Document saved as {file_path}")

create_cubing_report("C:/Users/Karthik/Downloads/Cubing_Event_Report_11.docx", api_data)
